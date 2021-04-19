import docx
import os
from datetime import date as dt
from datetime import datetime
from docx.shared import Pt
from docx.shared import Mm


def editSBAR(doc, templist):
    temptxt = doc.tables[0].columns[templist[2]].cells[templist[3]].paragraphs[0].add_run(templist[0])
    temptxt.font.name = 'Calibri'
    temptxt.font.size = Pt(11)
    if templist[1] == 'bold':
        temptxt.bold = True
    else:
        temptxt.bold = False
    return doc


def saveSBAR(report_doc, name):
    section = report_doc.sections[0]
    section.page_height = Mm(210)
    section.page_width = Mm(297)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.top_margin = Mm(55)
    section.bottom_margin = Mm(10)
    section.header_distance = Mm(28)
    section.footer_distance = Mm(12)
    filenameout = 'C:/PyCharmProjects/assets/report_%s_%s.docx' % name
    report_doc.save(filenameout)


def completeReport(sbar, setting):
    doc = docx.Document('C:/PyCharmProjects/assets/communityReport.docx')
    os.getcwd()
    today = dt.today()
    today_str = today.strftime('%d %B, %Y')
    now = datetime.now()
    time_str = now.strftime("%H:%M")

    if setting == 'community':

        # Situation - grid locations columns 0,1 rows 2,7 [8/9 for date/clinician] - [8 rows]

        temp = [sbar.situation0, 'normal', 0, 3]
        doc = editSBAR(doc, temp)

        temp = [sbar.situation1, 'normal', 0, 4]
        doc = editSBAR(doc, temp)

        temp = ['Patient Details', 'bold', 0, 6]
        doc = editSBAR(doc, temp)

        temp = [sbar.background0, 'normal', 0, 7]
        doc = editSBAR(doc, temp)

        # tempstr = 'Mobile Telephone: %s' % sbar.mob
        # temp = [tempstr, 'normal', 0, 5]
        # doc = editSBAR(doc, temp)

        # Date
        temp = ['Assessment date', 'bold', 0, 8]
        doc = editSBAR(doc, temp)
        # Date string
        datestr = '%s, %s' % (today_str, time_str)
        temp = [datestr, 'normal', 1, 8]
        doc = editSBAR(doc, temp)
        # Assessor
        temp = ['Clinician', 'bold', 0, 9]
        doc = editSBAR(doc, temp)
        clinstr = 'Dr B Kemp, ST7 Bleep 595'
        temp = [clinstr, 'normal', 1, 8]
        doc = editSBAR(doc, temp)

        # doc.tables[0].columns[1].cells[1].text = today_str
        # # Time
        # doc.tables[0].columns[0].cells[2].text = 'Time'
        # doc.tables[0].columns[1].cells[2].text = time_str
        # # setting
        # doc.tables[0].columns[0].cells[3].text = 'Setting'
        # doc.tables[0].columns[1].cells[3].text = 'Community'
        # # staff name
        # doc.tables[0].columns[0].cells[4].text = 'Clinician'
        # doc.tables[0].columns[1].cells[4].text = 'B Kemp [Doctor ST7]'
        # # staff role
        # # doc.tables[0].columns[0].cells[5].text = 'Role'
        # # doc.tables[0].columns[1].cells[5].text = 'Doctor - ST7'
        #
        # # Situation
        # doc.tables[0].columns[0].cells[10].text = sbar.situation0
        # doc.tables[0].columns[0].cells[11].text = sbar.situation1

        # Background - column 0, rows 13 to 23 [11 rows]
        doc.tables[0].columns[1].cells[16].text = sbar.name
        doc.tables[0].columns[1].cells[17].text = sbar.nhs
        doc.tables[0].columns[1].cells[18].text = sbar.dob
        doc.tables[0].columns[1].cells[19].text = sbar.mob
        doc.tables[0].columns[0].cells[20].text = ''
        doc.tables[0].columns[0].cells[21].text = sbar.background1

        # Assessment column 2, rows 15 to 23 [9 rows]
        doc.tables[0].columns[2].cells[13].text = sbar.assess0
        doc.tables[0].columns[2].cells[16].text = sbar.assess1
        doc.tables[0].columns[2].cells[17].text = sbar.assess2
        doc.tables[0].columns[2].cells[18].text = sbar.assess3

        # Recommendation column 3, rows 15 to 23 [9 rows]
        doc.tables[0].columns[4].cells[12].text = sbar.response0
        doc.tables[0].columns[4].cells[13].text = sbar.response1
        doc.tables[0].columns[4].cells[14].text = sbar.response2
        doc.tables[0].columns[4].cells[15].text = sbar.response3
